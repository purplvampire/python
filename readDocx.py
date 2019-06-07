#! python3
from docx import Document
def getText(filename):
    doc = Document(filename)
    fullText = []
    for paraNum in range(len(doc.paragraphs)):
        para = doc.paragraphs[paraNum].text
        fullText.append(para)
        # print(fullText)
    return '\n'.join(fullText)      # Connect list to string by \n
